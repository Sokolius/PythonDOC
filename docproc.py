from docx import Document
from  docx.shared import Pt
from createpadate import FolderCreator
from datetime import datetime, timedelta
from docx.shared import Inches


class Documentproc:
    def __init__(self,t_path):
        folder = FolderCreator('D://_BIG_DATA/CurrencyMarketPython/AddeddatatobdPyton/DocHistory')
        self.pathfold = folder.create_folder()
        self.doc = Document(t_path)
        self.inline =[] # прогоны для редактирования текста в абзаце

    def replace_text(self,placeholder,replacetext):
        for parag in self.doc.paragraphs:
            if placeholder in parag.text:
                parag.text = parag.text.replace(placeholder,replacetext)
                self.inline = parag.runs
                # inline = parag.runs
                # for run in inline:
                #     if placeholder in run.text:
                #         run.text = run.text.replace(placeholder,replacetext)

    def replace_img(self,placeholder,replaceimgpath):
        for parag in self.doc.paragraphs:
            if placeholder in parag.text:
                parag.clear()
                run = parag.add_run()
                run.add_picture(replaceimgpath,width=Inches(6),height=Inches(3))
                self.inline = parag.runs


    def set_font(self,font_name, font_size):
        for run in self.inline:
            run.font.name = font_name
            run.font.size = Pt(font_size)


    def sevedoc(self,filename):
        save_path = self.pathfold + "/"+filename
        self.doc.save(save_path)
